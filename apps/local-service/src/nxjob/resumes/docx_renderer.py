from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt

from nxjob.schemas.core import TailoredExperienceSection, TailoredResumeContent


def render_resume_docx(content: TailoredResumeContent, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    title_run = title.add_run(content.candidate_name)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(14.5)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    if content.contact_line:
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_before = Pt(0)
        contact.paragraph_format.space_after = Pt(0)
        contact_run = contact.add_run(content.contact_line)
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(_contact_font_size(content.contact_line))
        contact_run.font.color.rgb = RGBColor(0, 0, 0)

    _add_section(document, "PROFESSIONAL SUMMARY", content.summary, bullet=False)
    if content.skills:
        _add_section(document, "CORE QUALIFICATIONS / TECHNICAL SKILLS", [", ".join(content.skills)], bullet=False)
    if content.experience_sections:
        _add_experience_section(document, content.experience_sections)
    else:
        _add_section(document, "PROFESSIONAL EXPERIENCE", content.experience_bullets, bullet=True)
    if content.education:
        _add_section(document, "EDUCATION", content.education, bullet=False)

    document.save(output_path)
    return output_path


def _add_section(document: Document, heading: str, lines: list[str], bullet: bool) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    _add_section_heading(paragraph, heading)
    for line in lines:
        item = document.add_paragraph(style=None)
        item.paragraph_format.space_before = Pt(0)
        item.paragraph_format.space_after = Pt(0)
        item.paragraph_format.line_spacing = 1
        if bullet:
            item.paragraph_format.left_indent = Inches(0.18)
            item.paragraph_format.first_line_indent = Inches(-0.12)
            prefix = item.add_run("· ")
            _format_run(prefix, size=9, bold=False)
        run = item.add_run(line)
        _format_run(run, size=9, bold=False)


def _add_experience_section(document: Document, sections: list[TailoredExperienceSection]) -> None:
    paragraph = document.add_paragraph()
    _add_section_heading(paragraph, "PROFESSIONAL EXPERIENCE")
    for section in sections:
        header = document.add_paragraph(style=None)
        header.paragraph_format.space_before = Pt(1)
        header.paragraph_format.space_after = Pt(0)
        header.paragraph_format.line_spacing = 1
        run = header.add_run(_experience_header_line(section))
        _format_run(run, size=9.3, bold=True)

        for line in section.bullets:
            item = document.add_paragraph(style=None)
            item.paragraph_format.space_before = Pt(0)
            item.paragraph_format.space_after = Pt(0)
            item.paragraph_format.line_spacing = 1
            item.paragraph_format.left_indent = Inches(0.18)
            item.paragraph_format.first_line_indent = Inches(-0.12)
            prefix = item.add_run("· ")
            _format_run(prefix, size=9, bold=False)
            run = item.add_run(line)
            _format_run(run, size=9, bold=False)


def _experience_header_line(section: TailoredExperienceSection) -> str:
    role = " | ".join(part for part in [section.company, section.location] if part.strip())
    title = " | ".join(part for part in [section.title, section.date_range] if part.strip())
    if role and title:
        return f"{role} | {title}"
    return role or title


def _add_section_heading(paragraph, heading: str) -> None:
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1
    _add_bottom_border(paragraph)
    run = paragraph.add_run(heading)
    _format_run(run, size=10.5, bold=True)


def _format_run(run, size: float, bold: bool) -> None:
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def _contact_font_size(contact_line: str) -> float:
    length = len(contact_line)
    if length > 118:
        return 7.8
    if length > 104:
        return 8.3
    return 9
