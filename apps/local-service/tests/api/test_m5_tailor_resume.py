from __future__ import annotations

import json
import sqlite3
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from fastapi.testclient import TestClient

from nxjob.ai.openai_compatible import AiProviderError
from nxjob.db.repositories import new_id, utc_now
from nxjob.main import create_app
from nxjob.schemas.core import TailoredExperienceSection, TailoredResumeContent
from nxjob.workflows.resume_tailor import (
    TailorDraft,
    _contains_year_range,
    _draft_from_ai_payload,
    _fit_content_to_one_page_budget,
    _fit_contact_line,
    estimate_layout_budget,
)


class FakeAiResponse:
    def __init__(self, payload: dict) -> None:
        self.payload = payload

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, traceback) -> None:
        return None

    def read(self) -> bytes:
        return json.dumps(self.payload).encode("utf-8")


def test_tailor_resume_generates_docx_and_resume_version(tmp_path, monkeypatch) -> None:
    db_path = tmp_path / "nxjob.sqlite3"
    output_dir = tmp_path / "generated"
    monkeypatch.setenv("NXJOB_DB_PATH", str(db_path))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(output_dir))

    with TestClient(create_app()) as client:
        job_id = _capture_job(
            client,
            "Backend Engineer role using Python, FastAPI, SQLite, automation, and API workflows.",
        )
        response = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_id": "master_default",
                "master_resume_bullets": [
                    {
                        "id": "bullet_python_api",
                        "text": "Built Python FastAPI services with SQLite-backed workflow automation.",
                        "tags": ["Python", "FastAPI", "SQLite", "automation"],
                    },
                    {
                        "id": "bullet_frontend",
                        "text": "Improved React popup UI for browser extension workflows.",
                        "tags": ["React", "TypeScript"],
                    },
                ],
                "constraints": {
                    "format": "docx",
                    "target_length": "one_page_preferred",
                    "ats_friendly": True,
                },
                "success_reference_limit": 3,
            },
        )

    body = response.json()
    assert response.status_code == 200
    assert body["trace_id"].startswith("trc_")
    assert body["resume_version"]["format"] == "docx"
    assert body["resume_version"]["selected_bullets"][0] == "bullet_python_api"
    assert "Selected" in body["resume_version"]["change_summary"]
    assert body["docx_path"] == body["resume_version"]["file_path"]
    assert body["markdown_path"].endswith(".md")
    assert "_resume_20" in body["filename_base"]
    assert body["layout_budget"]["max_body_lines"] == 55
    assert body["quality_checks"]["summary_avoids_fixed_year_count"] is True

    file_path = Path(body["resume_version"]["file_path"])
    markdown_path = Path(body["markdown_path"])
    assert file_path.exists()
    assert markdown_path.exists()
    assert "Python FastAPI services" in markdown_path.read_text(encoding="utf-8")
    paragraphs = [paragraph.text for paragraph in Document(file_path).paragraphs]
    assert any("Python FastAPI services" in text for text in paragraphs)
    with ZipFile(file_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "w:noWrap" not in document_xml

    with sqlite3.connect(db_path) as connection:
        prompt_row = connection.execute(
            "SELECT provider, token_usage_json FROM prompt_logs WHERE trace_id = ?",
            (body["trace_id"],),
        ).fetchone()
        resume_row = connection.execute(
            "SELECT status FROM job_leads WHERE id = ?",
            (job_id,),
        ).fetchone()

    assert prompt_row[0] == "local_stub"
    assert json.loads(prompt_row[1])["input_chars"] > 0
    assert resume_row[0] == "tailored"


def test_tailor_resume_accepts_en_dash_education_years() -> None:
    assert _contains_year_range("Northeastern University, Boston, MA, 2018 – 2020")


def test_tailor_resume_contact_line_keeps_core_contact_on_one_line() -> None:
    contact = _fit_contact_line(
        "Boston / Greater Boston, MA | 555-010-0000 | candidate@example.com | "
        "LinkedIn: candidate-profile | H-1B transfer candidate, open to U.S. relocation"
    )

    assert contact == "Boston / Greater Boston, MA | 555-010-0000 | candidate@example.com | LinkedIn: candidate-profile"
    assert "H-1B" not in contact


def test_tailor_resume_compacts_ai_output_to_one_page_budget() -> None:
    content = TailoredResumeContent(
        candidate_name="Candidate",
        contact_line="Boston, MA | candidate@example.com",
        headline="Sustainability Analyst",
        summary=[
            "Analytical sustainability professional focused on energy, emissions, and reporting.",
            "Experienced in dashboards, financial modeling, and cross-functional implementation.",
        ],
        skills=[f"Skill {index}" for index in range(30)],
        experience_sections=[
            TailoredExperienceSection(
                company=f"Company {section}",
                location="Boston, MA",
                title="Analyst",
                date_range="2020 - Present",
                bullets=[
                    f"Built detailed analytical workflow and reporting artifact number {item} for energy, emissions, and operational decision support."
                    for item in range(8)
                ],
            )
            for section in range(6)
        ],
        experience_bullets=[],
        education=[
            "M.S. in Energy Systems Engineering, Northeastern University, Boston, MA, 2018 – 2020",
            "B.Eng. in Thermal Energy and Power Engineering, China University of Petroleum, Beijing, China, 2011 – 2015",
        ],
    )

    adjusted, warnings = _fit_content_to_one_page_budget(content)

    assert len(adjusted.skills) <= 18
    assert estimate_layout_budget(adjusted)["body_lines"] <= 55
    assert any("Compressed" in warning for warning in warnings)


def test_tailor_resume_adds_baseline_bullets_when_ai_output_underfills_page() -> None:
    content = TailoredResumeContent(
        candidate_name="Candidate",
        contact_line="Boston, MA | candidate@example.com",
        headline="Application Engineer",
        summary=["Application engineer focused on technical customer support."],
        skills=["Application Engineering", "Technical Sales Support"],
        experience_sections=[
            TailoredExperienceSection(
                company="BostonRen LLC",
                location="Boston, MA",
                title="Energy Analyst",
                date_range="2023 - Present",
                bullets=["Supported HVAC analysis and customer-facing technical deliverables."],
            )
        ],
        experience_bullets=[],
        education=["Northeastern University | M.S. | Boston, MA | 2018 - 2020"],
    )
    baseline_content = content.model_copy(
        update={
            "experience_sections": [
                TailoredExperienceSection(
                    company="BostonRen LLC",
                    location="Boston, MA",
                    title="Energy Analyst",
                    date_range="2023 - Present",
                    bullets=[
                        "Supported HVAC analysis and customer-facing technical deliverables.",
                        "Prepared bid-stage scopes, project-cost assumptions, and client-ready decision materials for engineered building upgrades.",
                        "Translated field findings, equipment constraints, and utility data into implementation-ready recommendations for stakeholders.",
                    ],
                )
            ]
        }
    )
    baseline = TailorDraft(
        content=baseline_content,
        selected_bullet_ids=[],
        change_summary="",
        token_usage={},
        markdown="",
        layout_budget=estimate_layout_budget(baseline_content),
        quality_checks={},
        warnings=[],
    )

    adjusted, warnings = _fit_content_to_one_page_budget(content, baseline)

    assert len(adjusted.experience_sections[0].bullets) == 3
    assert estimate_layout_budget(adjusted)["body_lines"] <= 55
    assert any("Added 2 truthful baseline bullet" in warning for warning in warnings)


def test_ai_payload_empty_content_is_rejected() -> None:
    baseline_content = TailoredResumeContent(
        candidate_name="Xu (Sonny) Shen",
        contact_line="Boston, MA | 555-010-0000 | candidate@example.com | LinkedIn: candidate-profile",
        headline="Application Engineer alignment",
        summary=["Engineer focused on energy systems, technical analysis, and customer-facing execution."],
        skills=["Application Engineering", "Technical Analysis", "Project Coordination"],
        experience_sections=[
            TailoredExperienceSection(
                company="BostonRen LLC",
                location="Boston, MA",
                title="Energy Analyst",
                date_range="2023 - Present",
                bullets=[
                    "Developed technical scopes and analysis for building-system upgrades and client decisions.",
                    "Supported project bids, decision decks, and cross-functional technical delivery.",
                ],
            )
        ],
        education=["Northeastern University | M.S. in Energy Systems Engineering | 2018 - 2020"],
    )
    baseline = TailorDraft(
        content=baseline_content,
        selected_bullet_ids=["baseline_bullet"],
        change_summary="Local baseline.",
        token_usage={},
        markdown="",
        layout_budget=estimate_layout_budget(baseline_content),
        quality_checks={},
        warnings=[],
    )

    with pytest.raises(AiProviderError) as exc_info:
        _draft_from_ai_payload({"content": {}}, baseline, {"total_tokens": 10})

    assert exc_info.value.category == "invalid_response"


def test_ai_payload_partial_content_is_repaired_from_local_baseline() -> None:
    baseline_content = TailoredResumeContent(
        candidate_name="Xu (Sonny) Shen",
        contact_line="Boston, MA | 555-010-0000 | candidate@example.com | LinkedIn: candidate-profile",
        headline="Application Engineer alignment",
        summary=["Engineer focused on energy systems, technical analysis, and customer-facing execution."],
        skills=["Application Engineering", "Technical Analysis", "Project Coordination"],
        experience_sections=[
            TailoredExperienceSection(
                company="BostonRen LLC",
                location="Boston, MA",
                title="Energy Analyst",
                date_range="2023 - Present",
                bullets=[
                    "Developed technical scopes and analysis for building-system upgrades and client decisions.",
                    "Supported project bids, decision decks, and cross-functional technical delivery.",
                ],
            )
        ],
        education=["Northeastern University | M.S. in Energy Systems Engineering | 2018 - 2020"],
    )
    baseline = TailorDraft(
        content=baseline_content,
        selected_bullet_ids=["baseline_bullet"],
        change_summary="Local baseline.",
        token_usage={},
        markdown="",
        layout_budget=estimate_layout_budget(baseline_content),
        quality_checks={},
        warnings=[],
    )

    draft = _draft_from_ai_payload(
        {
            "content": {
                "summary": ["Application engineer aligned with customer-facing technical support."],
                "quality_checks": {"truthful_to_master_resume": True},
            }
        },
        baseline,
        {"total_tokens": 10},
    )

    assert draft.content.candidate_name == "Xu (Sonny) Shen"
    assert draft.content.summary
    assert draft.content.skills
    assert draft.content.experience_sections
    assert draft.content.education
    assert "BostonRen LLC" in draft.markdown
    assert draft.quality_checks["ai_repaired_from_baseline"] is True
    assert draft.quality_checks["requires_user_review"] is True
    assert any("AI output was incomplete" in warning for warning in draft.warnings)


def test_ai_empty_content_does_not_generate_candidate_only_docx(tmp_path, monkeypatch) -> None:
    db_path = tmp_path / "nxjob.sqlite3"
    output_dir = tmp_path / "generated"
    master_path = tmp_path / "master-resume.json"
    master_path.write_text(
        json.dumps(
            {
                "id": "master_default",
                "candidate_name": "Xu (Sonny) Shen",
                "contact_line": "Boston, MA | 555-010-0000 | candidate@example.com",
                "bullets": [
                    {
                        "id": "bullet_energy",
                        "text": "Built energy planning analyses and client-ready technical recommendations.",
                        "tags": ["energy", "planning", "analysis"],
                    }
                ],
                "experience": [
                    {
                        "company": "BostonRen LLC",
                        "location": "Boston, MA",
                        "title": "Energy Analyst",
                        "start_date": "2023",
                        "end_date": "Present",
                        "bullets": [
                            {
                                "id": "exp_energy",
                                "text": "Developed technical scopes for energy-system upgrades and implementation planning.",
                                "tags": ["energy", "technical", "planning"],
                            }
                        ],
                    }
                ],
                "education": [
                    {
                        "school": "Northeastern University",
                        "degree": "M.S. in Energy Systems Engineering",
                        "location": "Boston, MA",
                        "start_year": "2018",
                        "end_year": "2020",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setenv("NXJOB_DB_PATH", str(db_path))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(output_dir))
    monkeypatch.setenv("NXJOB_MASTER_RESUME_PATH", str(master_path))
    monkeypatch.setenv("NXJOB_AI_API_KEY", "private-test-api-key")
    monkeypatch.setenv("NXJOB_AI_MODEL", "test-tailor-model")

    def fake_urlopen(request, timeout):
        return FakeAiResponse(
            {
                "model": "test-tailor-model",
                "usage": {"total_tokens": 42},
                "choices": [
                    {
                        "message": {
                            "content": json.dumps(
                                {
                                    "content": {},
                                    "selected_bullet_ids": [],
                                    "change_summary": "Provider returned no resume content.",
                                }
                            )
                        }
                    }
                ],
            }
        )

    monkeypatch.setattr("nxjob.ai.openai_compatible.urlopen", fake_urlopen)

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Energy planning role using technical analysis and stakeholder coordination.")
        response = client.post("/api/v1/resumes/tailor", json={"job_lead_id": job_id})

    assert response.status_code == 502
    assert response.json()["detail"]["error"]["code"] == "invalid_response"
    assert not output_dir.exists() or not list(output_dir.iterdir())

    with sqlite3.connect(db_path) as connection:
        resume_count = connection.execute("SELECT COUNT(*) FROM resume_versions").fetchone()[0]
        job_status = connection.execute("SELECT status FROM job_leads WHERE id = ?", (job_id,)).fetchone()[0]
        prompt_row = connection.execute("SELECT error FROM prompt_logs").fetchone()

    assert resume_count == 0
    assert job_status != "tailored"
    assert prompt_row[0] == "invalid_response"


def test_tailor_resume_uses_configured_ai_provider_without_logging_private_inputs(
    tmp_path,
    monkeypatch,
) -> None:
    db_path = tmp_path / "nxjob.sqlite3"
    monkeypatch.setenv("NXJOB_DB_PATH", str(db_path))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))
    monkeypatch.setenv("NXJOB_AI_API_KEY", "private-test-api-key")
    monkeypatch.setenv("NXJOB_AI_MODEL", "test-tailor-model")

    def fake_ai_tailor(*args, **kwargs):
        ai_config = args[3]
        assert ai_config.api_key == "private-test-api-key"
        content = TailoredResumeContent(
            candidate_name="Candidate",
            contact_line="candidate@example.com",
            headline="Backend Engineer alignment",
            summary=["Technical professional focused on Python API workflows."],
            skills=["Python", "FastAPI", "SQLite"],
            experience_bullets=["Built Python FastAPI services with SQLite workflow automation."],
            education=["Northeastern University | M.S. | Boston, MA | 2018 - 2020 | GPA: 3.62 / 4.0"],
        )
        return TailorDraft(
            content=content,
            selected_bullet_ids=["bullet_python_api"],
            change_summary="Generated structured resume content with AI provider.",
            token_usage={"prompt_tokens": 100, "completion_tokens": 80, "total_tokens": 180},
            markdown="# Candidate\n",
            layout_budget={"body_lines": 8, "heading_lines": 4, "max_body_lines": 55},
            quality_checks={"one_page_budget_ok": True},
            warnings=[],
        )

    monkeypatch.setattr("nxjob.api.resumes.tailor_resume_content_with_ai", fake_ai_tailor)

    with TestClient(create_app()) as client:
        job_id = _capture_job(
            client,
            "Private JD text for Backend Engineer role using Python APIs.",
        )
        response = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_bullets": [
                    {
                        "id": "bullet_python_api",
                        "text": "Private master resume bullet about Python FastAPI services.",
                        "tags": ["Python", "FastAPI"],
                    }
                ],
            },
        )

    assert response.status_code == 200
    body = response.json()
    assert body["ai_used"] is True
    assert body["ai_provider_name"] == "openai"

    with sqlite3.connect(db_path) as connection:
        prompt_row = connection.execute(
            """
            SELECT provider, model, input_summary, output_summary, token_usage_json, error
            FROM prompt_logs
            WHERE trace_id = ?
            """,
            (body["trace_id"],),
        ).fetchone()

    serialized_prompt_row = " ".join(str(value) for value in prompt_row)
    assert prompt_row[0] == "openai"
    assert prompt_row[1] == "test-tailor-model"
    assert json.loads(prompt_row[4])["total_tokens"] == 180
    assert "private-test-api-key" not in serialized_prompt_row
    assert "Private JD text" not in serialized_prompt_row
    assert "Private master resume bullet" not in serialized_prompt_row


def test_tailor_resume_ai_provider_failure_is_sanitized(tmp_path, monkeypatch) -> None:
    db_path = tmp_path / "nxjob.sqlite3"
    monkeypatch.setenv("NXJOB_DB_PATH", str(db_path))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))
    monkeypatch.setenv("NXJOB_AI_API_KEY", "private-test-api-key")
    monkeypatch.setenv("NXJOB_AI_MODEL", "test-tailor-model")

    def fake_ai_tailor(*args, **kwargs):
        raise AiProviderError("authentication_failed", "AI provider authentication failed.", 401)

    monkeypatch.setattr("nxjob.api.resumes.tailor_resume_content_with_ai", fake_ai_tailor)

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Private JD text for Python automation role.")
        response = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_bullets": [
                    {
                        "id": "bullet_api",
                        "text": "Private master resume bullet about API automation.",
                        "tags": ["Python", "API"],
                    }
                ],
            },
        )

    assert response.status_code == 401
    assert response.json()["detail"]["message"] == "AI provider authentication failed."
    assert response.json()["detail"]["error"]["code"] == "authentication_failed"
    assert response.json()["detail"]["error"]["config_source"] == "environment"
    assert "private-test-api-key" not in response.text

    with sqlite3.connect(db_path) as connection:
        prompt_row = connection.execute(
            "SELECT input_summary, output_summary, error FROM prompt_logs",
        ).fetchone()
        trace_row = connection.execute("SELECT status FROM workflow_traces").fetchone()

    serialized_prompt_row = " ".join(str(value) for value in prompt_row)
    assert prompt_row[2] == "authentication_failed"
    assert trace_row[0] == "failed"
    assert "private-test-api-key" not in serialized_prompt_row
    assert "Private JD text" not in serialized_prompt_row


def test_tailor_resume_reuses_cache_and_can_force_refresh(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Data automation role using Python and APIs.")
        payload = {
            "job_lead_id": job_id,
            "master_resume_bullets": [
                {
                    "id": "bullet_api",
                    "text": "Automated API workflows with Python services.",
                    "tags": ["Python", "API"],
                }
            ],
        }
        first = client.post("/api/v1/resumes/tailor", json=payload)
        second = client.post("/api/v1/resumes/tailor", json=payload)
        refreshed = client.post(
            "/api/v1/resumes/tailor",
            json={**payload, "force_refresh": True},
        )

    assert first.status_code == 200
    assert second.status_code == 200
    assert refreshed.status_code == 200
    assert first.json()["cache"]["hit"] is False
    assert second.json()["cache"]["hit"] is True
    assert refreshed.json()["cache"]["hit"] is False
    assert first.json()["resume_version"]["id"] == second.json()["resume_version"]["id"]
    assert refreshed.json()["resume_version"]["id"] != first.json()["resume_version"]["id"]
    assert refreshed.json()["filename_base"].endswith("_v2")


def test_tailor_resume_uses_success_references(tmp_path, monkeypatch) -> None:
    db_path = tmp_path / "nxjob.sqlite3"
    monkeypatch.setenv("NXJOB_DB_PATH", str(db_path))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Platform role with FastAPI and workflow automation.")
        _insert_success_reference(db_path, job_id, ["fastapi", "automation"])
        response = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_bullets": [
                    {
                        "id": "bullet_success_overlap",
                        "text": "Shipped FastAPI automation workflows for internal platforms.",
                        "tags": ["FastAPI", "automation"],
                    }
                ],
                "success_reference_limit": 2,
            },
        )

    assert response.status_code == 200
    assert response.json()["used_success_references"] == ["sref_test"]


def test_tailor_resume_includes_education_years_from_master_resume(tmp_path, monkeypatch) -> None:
    master_path = tmp_path / "master-resume.json"
    master_path.write_text(
        json.dumps(
            {
                "id": "master_default",
                "candidate_name": "Xu (Sonny) Shen",
                "contact_line": "Boston, MA | candidate@example.com",
                "bullets": [
                    {
                        "id": "bullet_python",
                        "text": "Automated Python API workflows for operational reporting.",
                        "tags": ["Python", "API", "automation"],
                    }
                ],
                "education": [
                    {
                        "school": "Northeastern University",
                        "degree": "M.S. in Energy Systems Engineering",
                        "location": "Boston, MA",
                        "start_year": "2018",
                        "end_year": "2020",
                        "gpa": "3.62 / 4.0",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))
    monkeypatch.setenv("NXJOB_MASTER_RESUME_PATH", str(master_path))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Python automation role for operational APIs.")
        response = client.post("/api/v1/resumes/tailor", json={"job_lead_id": job_id})

    assert response.status_code == 200
    body = response.json()
    assert body["quality_checks"]["education_years_present"] is True
    markdown = Path(body["markdown_path"]).read_text(encoding="utf-8")
    assert "2018 - 2020" in markdown
    assert "3.62 / 4.0" in markdown


def test_tailor_resume_preserves_structured_experience_timeline(tmp_path, monkeypatch) -> None:
    master_path = tmp_path / "master-resume.json"
    master_path.write_text(
        json.dumps(
            {
                "id": "master_default",
                "candidate_name": "Xu (Sonny) Shen",
                "contact_line": "Boston, MA | candidate@example.com",
                "experience": [
                    {
                        "company": "BostonRen LLC",
                        "location": "Boston, MA",
                        "title": "Operations Analyst",
                        "start_date": "2024",
                        "end_date": "Present",
                        "bullets": [
                            {
                                "id": "exp_current_api",
                                "text": "Automated Python API workflows for operations reporting and issue tracking.",
                                "tags": ["Python", "API", "operations"],
                            }
                        ],
                    },
                    {
                        "company": "Earlier Energy Co",
                        "location": "Shanghai, China",
                        "title": "Project Engineer",
                        "start_date": "2020",
                        "end_date": "2024",
                        "bullets": [
                            {
                                "id": "exp_prior_controls",
                                "text": "Supported field engineering coordination for facility controls projects.",
                                "tags": ["engineering", "controls"],
                            }
                        ],
                    },
                ],
                "education": [
                    {
                        "school": "Northeastern University",
                        "degree": "M.S. in Energy Systems Engineering",
                        "location": "Boston, MA",
                        "start_year": "2018",
                        "end_year": "2020",
                        "gpa": "3.62 / 4.0",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))
    monkeypatch.setenv("NXJOB_MASTER_RESUME_PATH", str(master_path))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Python API operations analyst role.")
        response = client.post("/api/v1/resumes/tailor", json={"job_lead_id": job_id})

    assert response.status_code == 200
    body = response.json()
    assert body["quality_checks"]["experience_timeline_preserved"] is True
    markdown = Path(body["markdown_path"]).read_text(encoding="utf-8")
    assert "BostonRen LLC | Boston, MA | Operations Analyst | 2024 - Present" in markdown
    assert "Earlier Energy Co | Shanghai, China | Project Engineer | 2020 - 2024" in markdown
    paragraphs = [paragraph.text for paragraph in Document(body["docx_path"]).paragraphs]
    assert any("BostonRen LLC" in paragraph and "2024 - Present" in paragraph for paragraph in paragraphs)
    assert any("Earlier Energy Co" in paragraph and "2020 - 2024" in paragraph for paragraph in paragraphs)


def test_tailor_resume_requires_configured_output_directory(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.delenv("NXJOB_GENERATED_RESUME_DIR", raising=False)
    monkeypatch.setenv("LOCALAPPDATA", str(tmp_path / "appdata"))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Backend automation role using Python APIs.")
        response = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_bullets": [
                    {
                        "id": "bullet_api",
                        "text": "Automated API workflows with Python services.",
                        "tags": ["Python", "API"],
                    }
                ],
            },
        )

    assert response.status_code == 422
    assert response.json()["detail"] == "Resume output folder is not configured."


def test_tailor_resume_uses_safe_date_company_job_filename(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))

    with TestClient(create_app()) as client:
        captured = client.post(
            "/api/v1/job-leads/capture",
            json={
                "source_url": "https://example.com/jobs/tailor-test",
                "source_site": "company_ats",
                "page_title": "Data/Automation: Engineer* at ACME|Controls?",
                "company_name": "ACME|Controls?",
                "job_title": "Data/Automation: Engineer*",
                "selected_text": "Data automation role using Python and APIs.",
            },
        )
        job_id = captured.json()["job_lead"]["id"]
        response = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "candidate_name": "Test Candidate",
                "master_resume_bullets": [
                    {
                        "id": "bullet_api",
                        "text": "Automated API workflows with Python services.",
                        "tags": ["Python", "API"],
                    }
                ],
            },
        )

    assert response.status_code == 200
    filename = response.json()["filename_base"]
    assert filename.startswith("Test_Candidate_ACME_Controls_Data_Automation_Engineer_resume_20")
    assert not any(character in filename for character in '/\\:*?"<>|')


def test_tailor_resume_feedback_is_saved(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Backend automation role using Python APIs.")
        tailored = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_bullets": [
                    {
                        "id": "bullet_api",
                        "text": "Automated API workflows with Python services.",
                        "tags": ["Python", "API"],
                    }
                ],
            },
        )
        response = client.post(
            "/api/v1/resumes/feedback",
            json={
                "job_lead_id": job_id,
                "resume_version_id": tailored.json()["resume_version"]["id"],
                "rating": "good_fit",
                "user_notes": "Strong enough for MVP.",
            },
        )

    assert response.status_code == 200
    assert response.json()["feedback"]["id"].startswith("rfb_")
    assert response.json()["feedback"]["rating"] == "good_fit"
    assert response.json()["feedback"]["candidate_status"] == ""


def test_tailor_resume_feedback_marks_success_candidate_status(tmp_path, monkeypatch) -> None:
    monkeypatch.setenv("NXJOB_DB_PATH", str(tmp_path / "nxjob.sqlite3"))
    monkeypatch.setenv("NXJOB_GENERATED_RESUME_DIR", str(tmp_path / "generated"))

    with TestClient(create_app()) as client:
        job_id = _capture_job(client, "Backend automation role using Python APIs.")
        tailored = client.post(
            "/api/v1/resumes/tailor",
            json={
                "job_lead_id": job_id,
                "master_resume_bullets": [
                    {
                        "id": "bullet_api",
                        "text": "Automated API workflows with Python services.",
                        "tags": ["Python", "API"],
                    }
                ],
            },
        )
        response = client.post(
            "/api/v1/resumes/feedback",
            json={
                "job_lead_id": job_id,
                "resume_version_id": tailored.json()["resume_version"]["id"],
                "rating": "save_success_candidate",
                "user_notes": "Possible future reference if this gets a screen.",
            },
        )

    assert response.status_code == 200
    body = response.json()
    assert body["feedback"]["rating"] == "save_success_candidate"
    assert body["feedback"]["candidate_status"] == "saved_as_success_reference_candidate"
    assert "Possible future reference" in body["feedback"]["user_notes"]


def _capture_job(client: TestClient, jd_text: str) -> str:
    response = client.post(
        "/api/v1/job-leads/capture",
        json={
            "source_url": "https://example.com/jobs/tailor-test",
            "source_site": "company_ats",
            "page_title": "Tailor test",
            "selected_text": jd_text,
        },
    )
    assert response.status_code == 200
    return response.json()["job_lead"]["id"]


def _insert_success_reference(db_path: Path, job_id: str, keywords: list[str]) -> None:
    resume_id = new_id("res")
    with sqlite3.connect(db_path) as connection:
        connection.execute(
            """
            INSERT INTO resume_versions (
              id, job_lead_id, source_master_resume_id, created_at, format, file_path
            )
            VALUES (?, ?, 'master_default', ?, 'docx', ?)
            """,
            (resume_id, job_id, utc_now(), "D:/tmp/success.docx"),
        )
        connection.execute(
            """
            INSERT INTO success_references (
              id, job_lead_id, resume_version_id, outcome_type, outcome_at, source,
              effective_keywords_json, effective_bullets_json
            )
            VALUES ('sref_test', ?, ?, 'screen', ?, 'manual', ?, ?)
            """,
            (
                job_id,
                resume_id,
                utc_now(),
                json.dumps(keywords),
                json.dumps(["bullet_success_overlap"]),
            ),
        )
        connection.commit()
